"""
File content extractor for different document formats
Supports PDF, DOC, DOCX, and TXT files
"""

import tempfile
import os
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document
import io

class FileContentExtractor:
    """Extract text content from various file formats"""
    
    @staticmethod
    def extract_text_from_file(file_content: bytes, filename: str) -> Optional[str]:
        """
        Extract text from uploaded file based on file extension
        
        Args:
            file_content: Raw file bytes
            filename: Original filename with extension
            
        Returns:
            Extracted text content or None if extraction fails
        """
        if not filename:
            return None
            
        file_ext = os.path.splitext(filename.lower())[1]
        
        try:
            if file_ext == '.pdf':
                return FileContentExtractor._extract_from_pdf(file_content)
            elif file_ext in ['.docx', '.doc']:
                return FileContentExtractor._extract_from_docx(file_content)
            elif file_ext in ['.txt', '.text']:
                return FileContentExtractor._extract_from_txt(file_content)
            else:
                # Try to decode as text for unknown formats
                return FileContentExtractor._extract_from_txt(file_content)
        except Exception as e:
            print(f"Error extracting text from {filename}: {e}")
            return None
    
    @staticmethod
    def _extract_from_pdf(file_content: bytes) -> Optional[str]:
        """Extract text from PDF file"""
        text_content = ""
        
        # Try pdfplumber first (better for complex PDFs)
        try:
            with io.BytesIO(file_content) as pdf_file:
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
            
            if text_content.strip():
                return text_content.strip()
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}")
        
        # Fallback to PyPDF2
        try:
            with io.BytesIO(file_content) as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text_content = ""
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n"
                
                return text_content.strip() if text_content.strip() else None
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
            return None
    
    @staticmethod
    def _extract_from_docx(file_content: bytes) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            with io.BytesIO(file_content) as docx_file:
                doc = Document(docx_file)
                text_content = ""
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + " "
                    text_content += "\n"
                
                return text_content.strip() if text_content.strip() else None
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
            return None
    
    @staticmethod
    def _extract_from_txt(file_content: bytes) -> Optional[str]:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    return text_content.strip() if text_content.strip() else None
                except UnicodeDecodeError:
                    continue
            
            return None
        except Exception as e:
            print(f"TXT extraction failed: {e}")
            return None

def extract_resume_text(file_content: bytes, filename: str) -> str:
    """
    Main function to extract text from resume file
    
    Args:
        file_content: Raw file bytes
        filename: Original filename
        
    Returns:
        Extracted text or empty string if extraction fails
    """
    extracted_text = FileContentExtractor.extract_text_from_file(file_content, filename)
    
    if extracted_text:
        return extracted_text
    else:
        print(f"Failed to extract text from {filename}")
        return ""

# Test function
if __name__ == "__main__":
    # Test with sample text
    sample_text = b"""
    John Doe
    Software Engineer
    Email: john.doe@example.com
    Phone: +1-234-567-8900
    
    Skills: Python, JavaScript, React, Node.js
    Experience: 3 years in web development
    Education: B.Tech Computer Science
    """
    
    result = extract_resume_text(sample_text, "test.txt")
    print("Extracted text:")
    print(result)